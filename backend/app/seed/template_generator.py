import os
import docx

def create_mock_template(output_dir: str):
    doc = docx.Document()
    
    # Title
    doc.add_heading('ĐƠN ĐĂNG KÝ THAM GIA CHƯƠNG TRÌNH HỖ TRỢ DOANH NGHIỆP', level=1)
    
    # Intro
    doc.add_paragraph('Kính gửi: Ban quản lý chương trình hỗ trợ phát triển công nghệ.')
    
    doc.add_paragraph('Doanh nghiệp chúng tôi đăng ký tham gia chương trình hỗ trợ với thông tin chi tiết dưới đây:')
    
    # Fields (Jinja placeholders)
    doc.add_paragraph('1. Tên doanh nghiệp: {{ company_name }}')
    doc.add_paragraph('2. Mã số thuế: {{ tax_code }}')
    doc.add_paragraph('3. Địa chỉ trụ sở chính: {{ location }}')
    doc.add_paragraph('4. Tổng số lao động: {{ employee_count }}')
    doc.add_paragraph('5. Tỷ lệ chi cho hoạt động R&D: {{ rd_spend_ratio }}')
    doc.add_paragraph('6. Vốn điều lệ đăng ký: {{ registered_capital }}')
    doc.add_paragraph('7. Tổng doanh thu năm gần nhất: {{ revenue }}')
    
    doc.add_paragraph('Chúng tôi cam kết các thông tin khai báo trên là hoàn toàn chính xác và chịu trách nhiệm trước pháp luật.')
    
    # Signatures
    doc.add_paragraph('Đại diện pháp luật của doanh nghiệp\n(Ký tên, đóng dấu)')
    
    # Save
    output_path = os.path.join(output_dir, "grant_template.docx")
    doc.save(output_path)
    print(f"Generated mock template .docx at: {output_path}")

if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    create_mock_template(current_dir)
